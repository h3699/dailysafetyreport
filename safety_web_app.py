import streamlit as st
from docx import Document
from docx.shared import Inches
from datetime import datetime
import matplotlib.pyplot as plt
import pandas as pd
from io import BytesIO
import os

st.set_page_config(page_title="廣華醫院2期地盤安全報告", layout="wide")

# Logo
st.markdown("""
<div style="text-align: center; margin-bottom: 10px;">
    <img src="https://raw.githubusercontent.com/h3699/dailysafetyreport/main/logo.png" width="280">
</div>
""", unsafe_allow_html=True)

st.title("中國水電 俊和")
st.subheader("廣華醫院2期地盤安全巡查報告系統")

company = "中國水電 俊和"
site_name = "廣華醫院2期地盤"

password = st.text_input("🔐 輸入安全密碼", type="password")
if not password or password != st.secrets["general"]["password"]:
    if password:
        st.error("❌ 密碼錯誤！")
    st.stop()

categories = ["地盤整潔", "機械", "個人防護", "吊運", "高空工作", "離地工作", "電力安全", "分判管理", "通道"]

if 'issues' not in st.session_state:
    st.session_state.issues = []
if 'map_image' not in st.session_state:
    st.session_state.map_image = None

tab1, tab2 = st.tabs(["新增問題", "生成報告"])

with tab1:
    st.subheader("新增安全問題")
    
    # 地盤分區地圖
    st.write("**上傳地盤分區地圖** (永久保存)")
    uploaded_map = st.file_uploader("上傳地圖", type=["jpg", "jpeg", "png"], key="map_key")
    if uploaded_map is not None:
        st.session_state.map_image = uploaded_map
    if st.session_state.get('map_image'):
        st.image(st.session_state.map_image, width=700, caption="目前使用的地盤分區地圖")
        if st.button("🗑️ 刪除地圖"):
            st.session_state.map_image = None
            st.rerun()
    
    col1, col2 = st.columns([1, 1])
    with col1:
        date = st.date_input("巡查日期", datetime.today())
        location = st.text_input("發生地點")
        subcontractor = st.text_input("分判")
        category = st.selectbox("問題分類", categories)
        severity = st.selectbox("嚴重度", ["高", "中", "低"])
    with col2:
        problem = st.text_area("問題事項描述")
        suggestion = st.text_area("建議處理方法")
    uploaded_files = st.file_uploader("上傳問題相片", type=["jpg", "jpeg", "png"], accept_multiple_files=True)
    
    if st.button("✅ 新增此問題", type="primary"):
        if problem.strip() and location.strip():
            st.session_state.issues.append({
                "date": date,
                "location": location,
                "subcontractor": subcontractor,
                "category": category,
                "severity": severity,
                "problem": problem,
                "suggestion": suggestion,
                "photos": uploaded_files
            })
            st.success("✅ 已成功新增！")
            st.rerun()
        else:
            st.error("請填寫發生地點和問題事項")
    
    # 當日新增問題顯示
    if st.session_state.issues:
        st.subheader("當日新增問題")
        for i, issue in enumerate(st.session_state.issues):
            with st.expander(f"問題 {i+1} | {issue['category']} | {issue['severity']} | {issue['location']}"):
                st.write(f"**分判**：{issue.get('subcontractor', '未填')}")
                st.write(f"**問題**：{issue['problem']}")
                st.write(f"**建議**：{issue['suggestion']}")
                if issue.get('photos'):
                    for p in issue['photos']:
                        st.image(p, width=500)
                if st.button(f"🗑️ 刪除此問題", key=f"del_{i}"):
                    st.session_state.issues.pop(i)
                    st.success("✅ 已刪除")
                    st.rerun()
                    
with tab2:
    if st.button("🚀 生成 Word 報告", type="primary"):
        if not st.session_state.issues:
            st.error("請先新增問題")
        else:
            doc = Document()
            doc.add_heading(f'{company}\n{site_name}\n安全巡查問題分析報告', 0)
            doc.add_paragraph(f'報告日期：{datetime.now().strftime("%Y年%m月%d日")}')
            
            for i, issue in enumerate(st.session_state.issues):
                doc.add_heading(f'問題 {i+1} - {issue["location"]}', level=1)
                doc.add_paragraph(f'分判：{issue["subcontractor"]}   分類：{issue["category"]}   嚴重度：{issue["severity"]}')
                doc.add_paragraph(f'問題事項：{issue["problem"]}')
                doc.add_paragraph(f'建議處理方法：{issue["suggestion"]}')
                if issue.get('photos'):
                    for photo in issue['photos']:
                        try:
                            doc.add_picture(BytesIO(photo.getvalue()), width=Inches(4.5))
                        except:
                            pass
                doc.add_paragraph("")
            
            # 圖表放在最後
            doc.add_page_break()
            doc.add_heading('問題統計圖表', level=1)
            cat_count = pd.Series([i['category'] for i in st.session_state.issues]).value_counts()
            fig, ax = plt.subplots(figsize=(10, 6))
            cat_count.plot(kind='bar', ax=ax, color='skyblue')
            ax.set_title('各類型問題統計')
            ax.set_ylabel('問題數量')
            plt.xticks(rotation=45)
            plt.tight_layout()
            plt.savefig("chart.png", dpi=200)
            doc.add_picture("chart.png", width=Inches(6))
            
            filename = f"廣華醫院2期安全報告_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
            doc.save(filename)
            
            with open(filename, "rb") as f:
                st.download_button("📥 下載報告", f, file_name=filename)

st.caption("中國水電 俊和 | 廣華醫院2期地盤安全工具")
